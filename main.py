import os
import subprocess
from datetime import datetime
from docx import Document

root_path = os.path.abspath(os.curdir)
folder_path = root_path + "/templates/"
output_dir = root_path + "/data/pdf/"


def main(
    user_name,
    job_title,
    company_name,
):
    template_file_path = folder_path + f"{job_title.lower()}_cover_letter.docx"
    output_file_path = root_path + "/data/" + f"{user_name}-cover_letter.docx"

    current_date = datetime.now().strftime("%d %B, %Y")

    variables = {
        "${DATE}": current_date,
        "${JOB_TITLE}": job_title + "Developer",
        "${COMPANY_NAME}": company_name,
        "${YOUR_NAME}": user_name,
    }
    template_document = Document(template_file_path)
    try:
        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(
                                paragraph, variable_key, variable_value
                            )

        template_document.save(output_file_path)
        doc2pdf(output_file_path)
    except Exception as e:
        raise e


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


try:
    from comtypes import client
except ImportError:
    client = None


def doc2pdf(doc):
    """
    convert a doc/docx document to pdf format
    :param doc: path to document
    """
    doc = os.path.abspath(doc)  # bugfix - searching files in windows/system32
    if client is None:
        return doc2pdf_linux(doc, output_dir)
    name, ext = os.path.splitext(doc)
    try:
        word = client.CreateObject("Word.Application")
        worddoc = word.Documents.Open(doc)
        worddoc.SaveAs(name + ".pdf", FileFormat=17)
    except Exception:
        raise
    finally:
        worddoc.Close()
        word.Quit()


def doc2pdf_linux(doc, output_path):
    """
    convert a doc/docx document to pdf format (linux only, requires libreoffice)
    :param doc: path to document
    """
    # cmd = "libreoffice --convert-to pdf --outdir".split() + [doc, output_path]
    # p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    # p.wait(timeout=10)
    # stdout, stderr = p.communicate()
    # if stderr:
    #     raise subprocess.SubprocessError(stderr)
    try:
        os.system(
            "libreoffice --convert-to pdf" + " " + doc + " --outdir " + output_path
        )
    except Exception as e:
        raise e


# define Python user-defined exceptions
class InvalidChoiceException(Exception):
    "Raised when the input value is more than listed options"
    pass


if __name__ == "__main__":

    options = ["Laravel", "Python", "FullStack"]

    user_input = ""

    input_message = "Pick an Cover Letter Template:\n"

    for index, item in enumerate(options):
        input_message += f"{index+1}) {item}\n"

    print(input_message)

    while True:
        try:
            job_title = int(input("Enter your cover letter: "))
            if not job_title:
                raise ValueError
            if job_title > len(options):
                raise InvalidChoiceException
        except ValueError:
            print("Empty choose is not Accepted")
            continue
        except InvalidChoiceException:
            print("Invalid option choose")
            continue
        else:
            break
    while True:
        try:
            user_name = str(input("Enter your Name: "))
            if not user_name:
                raise ValueError
        except ValueError:
            print("Empty Name is not Accepted")
            continue
        else:
            break
    while True:
        try:
            company_name = str(input("Enter your Company Name: "))
            if not company_name:
                raise ValueError
        except ValueError:
            print("Empty Company Name is not Accepted")
            continue
        else:
            break
    main(
        user_name,
        options[job_title - 1],
        company_name,
    )
